from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import requests, re, os, pythoncom

app = Flask(__name__, static_folder="public", static_url_path="/", template_folder="public")
CORS(app)

API_URL = "https://api-inference.huggingface.co/models/mistralai/Mixtral-8x7B-Instruct-v0.1"
headers = {"Authorization": "Bearer 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"}

def query_llm(prompt):
    payload = {
        "inputs": prompt,
        "parameters": {
            "temperature": 0.7,
            "max_new_tokens": 2048
        }
    }
    response = requests.post(API_URL, headers=headers, json=payload)
    
    try:
        result = response.json()
        print("Ответ от Hugging face", result)

        if isinstance(result, dict) and result.get("error"):
            raise Exception(f"Ошибка от модели: {result['error']}")

        if isinstance(result, list) and "generated_text" in result[0]:
            return result[0]["generated_text"]
        
        if isinstance(result, list) and "output" in result[0]:
            return result[0]["output"]

        raise Exception("Не удалось распознать структуру ответа от модели.")
    
    except Exception as e:
        print("Ошибка при получении ответа от ИИ:")
        print(e)
        print("Ответ API:")
        print(response.text)
        return ""


def parse_ai_response(text):
    parts = re.split(r'\n(?=\d+\.\s+\*\*)', text.strip())
    data = {}
    for part in parts:
        match = re.match(r'(\d+)\.\s+\*\*(.+?)\*\*\n(.+)', part, re.DOTALL)
        if match:
            _, title, content = match.groups()
            data[title.strip()] = content.strip()
    return data


def create_pdf_preview(pdf_path, preview_path="public/images/preview.jpg", dpi=200):
    images = convert_from_path(pdf_path, dpi=dpi)

    if images:
        preview_image = images[0]
        preview_image.save(preview_path, 'JPEG')

    return preview_path


@app.route('/generate-plan', methods=['POST'])
def get_input_data():
    pythoncom.CoInitialize()

    # if os.path.isfile("public/documents/План учебного занятия.pdf"):
    #     preview_path = create_pdf_preview("public/documents/План учебного занятия.pdf")
    #     return jsonify({
    #         "status": "success",
    #         "message": "План получен",
    #         "preview_path": "images/preview.jpg"
    #     })      

    # if os.path.isfile("public/documents/План учебного занятия.docx"):
    #     convert("public/documents/План учебного занятия.docx", "public/documents/План учебного занятия.pdf")
    #     preview_path = create_pdf_preview("public/documents/План учебного занятия.pdf")
    #     return jsonify({
    #         "status": "success",
    #         "message": "План получен",
    #         "preview_path": "images/preview.jpg"
    #     })      
    
    
    input_data = request.get_json()
    print("Получены данные:", input_data)

    prompt = f"""
Пожалуйста, составь поурочный план учебного занятия, заполнив все поля, приведённые ниже. План должен быть подробным, логично структурированным и соответствовать теме, типу занятия и предполагаемой аудитории.

Входные данные:
- Тема занятия: {input_data['Тема занятия']}
- Модуль / Дисциплина: {input_data['Модуль']}
- Специальность / Квалификация: {input_data['Специальность']}
- Курс / Группа: {input_data['Курс, группа']}
- Тип занятия: {input_data['Тип занятия']}
- Продолжительность: {input_data['Продолжительность']}
- Предпочитаемые виды деятельности: {input_data['Виды деятельности']}

Заполни следующие поля:
1. **Тема занятия**
2. **Курс, группа**
3. **Тип занятия**
4. **Цели** (Образовательная:, Развивающая:, Воспитательная:)
5. **Задачи**
6. **Ожидаемые результаты**
7. **Необходимые ресурсы**
8. **Ход занятия** (1. Актуализация пройденного материала: (можно в виде опроса, тестов, карточек), 2. Объяснение нового материала: (лекция, презентация), 3. Закрепление нового материала: (зависит от преподаваемого предмета (в виде тестов, карточек, индивидуальной работы на компьютере, групповое и т.д), 4. Задание на дом:)
9. **Технология обучения**
10. **Межпредметная связь**
11. **Рефлексия по содержанию**
12. **Литература**
13. **Оценивание обучающихся**
14. **Домашнее задание**

Формат: текст с чёткими заголовками для каждого пункта.
"""

    print("⏳ Генерация поурочного плана...")
    raw_response = query_llm(prompt)
    print("Полученный ответ:", raw_response)
    parsed_data = parse_ai_response(raw_response)
    print("Форматированный ответ:", parsed_data)

    template_path = "public/documents/Шаблон.docx"
    doc = Document(template_path)

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                first_cell_text = row.cells[0].text.strip()
                for key in parsed_data:
                    if key.lower() in first_cell_text.lower():
                        cell = row.cells[1]
                        cell.text = ""

                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(parsed_data[key])
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(14)

                        rPr = run._element.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rPr.append(rFonts)

    output_path = "public/documents/План учебного занятия.docx"
    doc.save(output_path)
    convert("public/documents/План учебного занятия.docx", "public/documents/План учебного занятия.pdf") 
    create_pdf_preview("public/documents/План учебного занятия.pdf")
    print(f"План успешно создан: {output_path}")

    return jsonify({
        "status": "success",
        "message": "План получен",
        "preview_path": "images/preview.jpg"
    }) 


@app.route('/download-pdf')
def download_pdf():
    path_to_pdf = 'public/documents/План учебного занятия.docx'
    return send_file(path_to_pdf, as_attachment=True)


@app.route("/")
def home():
    return app.send_static_file("index.html")

@app.route("/start")
def formPage():
    return app.send_static_file("form.html")

if __name__ == "__main__":
    # from waitress import serve
    # serve(app, host="0.0.0.0", port=8080)
    app.run(debug=True, port=5000)